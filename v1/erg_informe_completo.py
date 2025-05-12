def generar_informe_word_completo(data, nombre_raton, info_adicional, grupo_nombres, UMBRAL_TIEMPO=0.1):
    import matplotlib.pyplot as plt
    import numpy as np
    import io
    from docx import Document
    from docx.shared import Inches
    from scipy.signal import find_peaks
    import os

    doc = Document()
    doc.add_heading("Informe de señales ERG", 0)
    doc.add_paragraph(f"Nombre del ratón: {nombre_raton}")
    doc.add_paragraph(f"Observaciones: {info_adicional}")

    tiempo_total = data['Tiempo'].unique()
    tiempo_ms = tiempo_total * 1000

    # ====== ESCOTÓPICO ======
    doc.add_heading("1. Escotópico", level=1)
    resultados_escotopico = []
    fig_izq, ax_izq = plt.subplots(figsize=(6, 4))
    fig_der, ax_der = plt.subplots(figsize=(6, 4))
    for grupo in [1, 2, 3]:
        nombre = grupo_nombres[grupo]
        subgrupo = data[data['Grupo'] == grupo]
        tiempo = subgrupo['Tiempo'].values
        tiempo_ms = tiempo * 1000
        ojo_izq = subgrupo['OjoIzquierdo'].values
        ojo_der = subgrupo['OjoDerecho'].values
        idx_inicio = np.argmax(tiempo >= UMBRAL_TIEMPO)

        for señal, ax, ojo_nombre in [(ojo_izq, ax_izq, "Ojo izquierdo"), (ojo_der, ax_der, "Ojo derecho")]:
            señal_post = señal[idx_inicio:]
            pico_b_rel = np.argmax(señal_post)
            pico_b_idx = idx_inicio + pico_b_rel
            pico_b_valor = señal[pico_b_idx]
            señal_inicial = señal[idx_inicio:pico_b_idx]
            min_rel = np.argmin(señal_inicial)
            inicio_b_idx = idx_inicio + min_rel
            ax.plot(tiempo_ms, señal, label=nombre)
            ax.plot(tiempo_ms[inicio_b_idx], señal[inicio_b_idx], 'mo')
            ax.plot(tiempo_ms[pico_b_idx], señal[pico_b_idx], 'ro')

            resultados_escotopico.append({
                "Grupo": nombre,
                "Ojo": ojo_nombre,
                "Inicio onda b (ms)": f"{tiempo_ms[inicio_b_idx]:.1f}",
                "Pico onda b (ms)": f"{tiempo_ms[pico_b_idx]:.1f}",
                "Amplitud onda b (µV)": f"{pico_b_valor - señal[inicio_b_idx]:.2f}"
            })

    for ax, title in [(ax_izq, "Ojo izquierdo - Comparativa escotópico"), (ax_der, "Ojo derecho - Comparativa escotópico")]:
        ax.set_title(title)
        ax.set_xlabel("Tiempo (ms)")
        ax.set_ylabel("µV")
        ax.legend()
        ax.grid()

    path_izq = "escotopico_izq.png"
    path_der = "escotopico_der.png"
    fig_izq.savefig(path_izq)
    fig_der.savefig(path_der)
    plt.close(fig_izq)
    plt.close(fig_der)

    row = doc.add_table(rows=1, cols=2).rows[0].cells
    row[0].paragraphs[0].add_run().add_picture(path_izq, width=Inches(3.0))
    row[1].paragraphs[0].add_run().add_picture(path_der, width=Inches(3.0))

    doc.add_heading("Tabla resumen de Escotópico", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Grupo'
    hdr[1].text = 'Ojo'
    hdr[2].text = 'Inicio (ms)'
    hdr[3].text = 'Pico (ms)'
    hdr[4].text = 'Amplitud (µV)'
    for res in resultados_escotopico:
        row = table.add_row().cells
        row[0].text = res['Grupo']
        row[1].text = res['Ojo']
        row[2].text = res['Inicio onda b (ms)']
        row[3].text = res['Pico onda b (ms)']
        row[4].text = res['Amplitud onda b (µV)']

    # ====== MESÓPICO ======
    doc.add_page_break() #salto de pagina

    doc.add_heading("2. Mesópico", level=1)
    resultados_mesopico = []
    fig_mes_izq, ax_mes_izq = plt.subplots(figsize=(6, 4))
    fig_mes_der, ax_mes_der = plt.subplots(figsize=(6, 4))
    for grupo in [4, 5]:
        nombre = grupo_nombres[grupo]
        subgrupo = data[data['Grupo'] == grupo]
        tiempo = subgrupo['Tiempo'].values
        tiempo_ms = tiempo * 1000
        ojo_izq = subgrupo['OjoIzquierdo'].values
        ojo_der = subgrupo['OjoDerecho'].values
        idx_inicio = np.argmax(tiempo >= UMBRAL_TIEMPO)

        for señal, ax, ojo_nombre in [(ojo_izq, ax_mes_izq, "Ojo izquierdo"), (ojo_der, ax_mes_der, "Ojo derecho")]:
            señal_post = señal[idx_inicio:]
            pico_b_rel = np.argmax(señal_post)
            pico_b_idx = idx_inicio + pico_b_rel
            pico_b_valor = señal[pico_b_idx]
            señal_a = señal[idx_inicio:pico_b_idx]
            pico_a_rel = np.argmin(señal_a)
            pico_a_idx = idx_inicio + pico_a_rel
            delta_b = pico_b_valor - señal[pico_a_idx]
            delta_a = señal[idx_inicio] - señal[pico_a_idx]

            ax.plot(tiempo_ms, señal, label=nombre)
            ax.plot(tiempo_ms[pico_a_idx], señal[pico_a_idx], 'go')
            ax.plot(tiempo_ms[pico_b_idx], señal[pico_b_idx], 'ro')

            resultados_mesopico.append({
                "Grupo": nombre,
                "Ojo": ojo_nombre,
                "Pico onda a (ms)": f"{tiempo_ms[pico_a_idx]:.1f}",
                "Pico onda b (ms)": f"{tiempo_ms[pico_b_idx]:.1f}",
                "Amplitud a (µV)": f"{delta_a:.2f}",
                "Amplitud b (µV)": f"{delta_b:.2f}"
            })

    ax_mes_izq.set_title("Ojo izquierdo - Comparativa mesópico")
    ax_mes_izq.set_xlabel("Tiempo (ms)")
    ax_mes_izq.set_ylabel("µV")
    ax_mes_izq.legend()
    ax_mes_izq.grid()

    ax_mes_der.set_title("Ojo derecho - Comparativa mesópico")
    ax_mes_der.set_xlabel("Tiempo (ms)")
    ax_mes_der.set_ylabel("µV")
    ax_mes_der.legend()
    ax_mes_der.grid()

    path_mes_izq = "mesopico_izq.png"
    path_mes_der = "mesopico_der.png"
    fig_mes_izq.savefig(path_mes_izq)
    fig_mes_der.savefig(path_mes_der)
    plt.close(fig_mes_izq)
    plt.close(fig_mes_der)

    row = doc.add_table(rows=1, cols=2).rows[0].cells
    row[0].paragraphs[0].add_run().add_picture(path_mes_izq, width=Inches(3.0))
    row[1].paragraphs[0].add_run().add_picture(path_mes_der, width=Inches(3.0))

    doc.add_heading("Tabla resumen de Mesópico", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Grupo'
    hdr[1].text = 'Ojo'
    hdr[2].text = 'Pico onda a (ms)'
    hdr[3].text = 'Pico onda b (ms)'
    hdr[4].text = 'Amplitud a (µV)'
    hdr[5].text = 'Amplitud b (µV)'
    for res in resultados_mesopico:
        row = table.add_row().cells
        row[0].text = res['Grupo']
        row[1].text = res['Ojo']
        row[2].text = res['Pico onda a (ms)']
        row[3].text = res['Pico onda b (ms)']
        row[4].text = res['Amplitud a (µV)']
        row[5].text = res['Amplitud b (µV)']

    # ====== POTENCIALES OSCILATORIOS ======
    doc.add_heading("3. Potenciales Oscilatorios", level=1)
    resultados_oscilatorios = []
    fig_osci_izq, ax_osci_izq = plt.subplots(figsize=(6, 4))
    fig_osci_der, ax_osci_der = plt.subplots(figsize=(6, 4))
    grupo = 6
    nombre = grupo_nombres[grupo]
    subgrupo = data[data['Grupo'] == grupo]
    tiempo = subgrupo['Tiempo'].values
    tiempo_ms = tiempo * 1000
    ojo_izq = subgrupo['OjoIzquierdo'].values
    ojo_der = subgrupo['OjoDerecho'].values

    for señal, ax, ojo_nombre in [(ojo_izq, ax_osci_izq, "Ojo izquierdo"), (ojo_der, ax_osci_der, "Ojo derecho")]:
        idx_inicio = np.argmax(tiempo >= 0.11)
        señal_post = señal[idx_inicio:]
        picos, _ = find_peaks(señal_post)

        if len(picos) < 2:
            continue

        picos_valores = señal_post[picos]
        top2_indices = np.argsort(picos_valores)[-2:]
        top2_sorted = top2_indices[np.argsort(picos_valores[top2_indices])]
        segundo_pico_rel = picos[top2_sorted[0]]
        segundo_pico_idx = idx_inicio + segundo_pico_rel
        segundo_pico_val = señal[segundo_pico_idx]

        pre_pico = señal[idx_inicio:segundo_pico_idx]
        min_pre_pico_idx_rel = np.argmin(pre_pico)
        min_pre_pico_idx = idx_inicio + min_pre_pico_idx_rel
        min_pre_pico_val = señal[min_pre_pico_idx]

        delta = segundo_pico_val - min_pre_pico_val

        ax.plot(tiempo_ms, señal, label=nombre)
        ax.plot(tiempo_ms[min_pre_pico_idx], min_pre_pico_val, 'mo', label='Inicio')
        ax.plot(tiempo_ms[segundo_pico_idx], segundo_pico_val, 'ro', label='Pico 2º')

        resultados_oscilatorios.append({
            "Ojo": ojo_nombre,
            "Inicio (ms)": f"{tiempo_ms[min_pre_pico_idx]:.1f}",
            "Pico (ms)": f"{tiempo_ms[segundo_pico_idx]:.1f}",
            "Amplitud (µV)": f"{delta:.2f}"
        })

    ax_osci_izq.set_title("Ojo izquierdo - Potenciales Oscilatorios")
    ax_osci_izq.set_xlabel("Tiempo (ms)")
    ax_osci_izq.set_ylabel("µV")
    ax_osci_izq.legend()
    ax_osci_izq.grid()

    ax_osci_der.set_title("Ojo derecho - Potenciales Oscilatorios")
    ax_osci_der.set_xlabel("Tiempo (ms)")
    ax_osci_der.set_ylabel("µV")
    ax_osci_der.legend()
    ax_osci_der.grid()

    path_osci_izq = "oscilatorios_izq.png"
    path_osci_der = "oscilatorios_der.png"
    fig_osci_izq.savefig(path_osci_izq)
    fig_osci_der.savefig(path_osci_der)
    plt.close(fig_osci_izq)
    plt.close(fig_osci_der)

    row = doc.add_table(rows=1, cols=2).rows[0].cells
    row[0].paragraphs[0].add_run().add_picture(path_osci_izq, width=Inches(3.0))
    row[1].paragraphs[0].add_run().add_picture(path_osci_der, width=Inches(3.0))

    doc.add_heading("Tabla resumen de Potenciales Oscilatorios", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Ojo'
    hdr[1].text = 'Inicio (ms)'
    hdr[2].text = 'Pico (ms)'
    hdr[3].text = 'Amplitud (µV)'
    for res in resultados_oscilatorios:
        row = table.add_row().cells
        row[0].text = res['Ojo']
        row[1].text = res['Inicio (ms)']
        row[2].text = res['Pico (ms)']
        row[3].text = res['Amplitud (µV)']

    # ====== FOTÓPICO ======
    doc.add_heading("4. Fotópico", level=1)
    resultados_fotopico = []
    fig_foto_izq, ax_foto_izq = plt.subplots(figsize=(6, 4))
    fig_foto_der, ax_foto_der = plt.subplots(figsize=(6, 4))
    for grupo in [7, 8]:
        nombre = grupo_nombres[grupo]
        subgrupo = data[data['Grupo'] == grupo]
        tiempo = subgrupo['Tiempo'].values
        tiempo_ms = tiempo * 1000
        ojo_izq = subgrupo['OjoIzquierdo'].values
        ojo_der = subgrupo['OjoDerecho'].values
        idx_inicio = np.argmax(tiempo >= UMBRAL_TIEMPO)

        for señal, ax, ojo_nombre in [(ojo_izq, ax_foto_izq, "Ojo izquierdo"), (ojo_der, ax_foto_der, "Ojo derecho")]:
            señal_post = señal[idx_inicio:]
            pico_b_rel = np.argmax(señal_post)
            pico_b_idx = idx_inicio + pico_b_rel
            pico_b_valor = señal[pico_b_idx]
            señal_a = señal[idx_inicio:pico_b_idx]
            pico_a_rel = np.argmin(señal_a)
            pico_a_idx = idx_inicio + pico_a_rel
            delta_b = pico_b_valor - señal[pico_a_idx]
            delta_a = señal[idx_inicio] - señal[pico_a_idx]

            ax.plot(tiempo_ms, señal, label=nombre)
            ax.plot(tiempo_ms[pico_a_idx], señal[pico_a_idx], 'go')
            ax.plot(tiempo_ms[pico_b_idx], señal[pico_b_idx], 'ro')

            resultados_fotopico.append({
                "Grupo": nombre,
                "Ojo": ojo_nombre,
                "Pico onda a (ms)": f"{tiempo_ms[pico_a_idx]:.1f}",
                "Pico onda b (ms)": f"{tiempo_ms[pico_b_idx]:.1f}",
                "Amplitud a (µV)": f"{delta_a:.2f}",
                "Amplitud b (µV)": f"{delta_b:.2f}"
            })

    ax_foto_izq.set_title("Ojo izquierdo - Comparativa fotópico")
    ax_foto_izq.set_xlabel("Tiempo (ms)")
    ax_foto_izq.set_ylabel("µV")
    ax_foto_izq.legend()
    ax_foto_izq.grid()

    ax_foto_der.set_title("Ojo derecho - Comparativa fotópico")
    ax_foto_der.set_xlabel("Tiempo (ms)")
    ax_foto_der.set_ylabel("µV")
    ax_foto_der.legend()
    ax_foto_der.grid()

    path_foto_izq = "fotopico_izq.png"
    path_foto_der = "fotopico_der.png"
    fig_foto_izq.savefig(path_foto_izq)
    fig_foto_der.savefig(path_foto_der)
    plt.close(fig_foto_izq)
    plt.close(fig_foto_der)

    row = doc.add_table(rows=1, cols=2).rows[0].cells
    row[0].paragraphs[0].add_run().add_picture(path_foto_izq, width=Inches(3.0))
    row[1].paragraphs[0].add_run().add_picture(path_foto_der, width=Inches(3.0))

    doc.add_heading("Tabla resumen de Fotópico", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Grupo'
    hdr[1].text = 'Ojo'
    hdr[2].text = 'Pico onda a (ms)'
    hdr[3].text = 'Pico onda b (ms)'
    hdr[4].text = 'Amplitud a (µV)'
    hdr[5].text = 'Amplitud b (µV)'
    for res in resultados_fotopico:
        row = table.add_row().cells
        row[0].text = res['Grupo']
        row[1].text = res['Ojo']
        row[2].text = res['Pico onda a (ms)']
        row[3].text = res['Pico onda b (ms)']
        row[4].text = res['Amplitud a (µV)']
        row[5].text = res['Amplitud b (µV)']

    
    doc.add_page_break() #salto de pagina

    # ====== FLICKER ======
    doc.add_heading("5. Flicker", level=1)
    resultados_flicker = []
    fig_flicker_izq, ax_flicker_izq = plt.subplots(figsize=(6, 4))
    fig_flicker_der, ax_flicker_der = plt.subplots(figsize=(6, 4))
    grupo = 9
    nombre = grupo_nombres[grupo]
    subgrupo = data[data['Grupo'] == grupo]
    tiempo = subgrupo['Tiempo'].values
    tiempo_ms = tiempo * 1000
    ojo_izq = subgrupo['OjoIzquierdo'].values
    ojo_der = subgrupo['OjoDerecho'].values

    for señal, ax, ojo_nombre in [(ojo_izq, ax_flicker_izq, "Ojo izquierdo"), (ojo_der, ax_flicker_der, "Ojo derecho")]:
        idx_inicio = np.argmax(tiempo >= 0.2)
        señal_post = señal[idx_inicio:]
        tiempo_post = tiempo_ms[idx_inicio:]
        picos_pos, _ = find_peaks(señal_post, prominence=20)
        picos_neg, _ = find_peaks(-señal_post, prominence=20)

        picos_total = sorted([(i, 'neg', señal_post[i]) for i in picos_neg] + [(i, 'pos', señal_post[i]) for i in picos_pos], key=lambda x: x[0])
        seleccionados = []
        for x in picos_total:
            if not seleccionados:
                if x[1] == 'neg':
                    seleccionados.append(x)
            elif len(seleccionados) % 2 == 1 and x[1] == 'pos':
                seleccionados.append(x)
            elif len(seleccionados) % 2 == 0 and x[1] == 'neg':
                seleccionados.append(x)
            if len(seleccionados) == 4:
                break

        for i, (idx_rel, tipo, valor) in enumerate(seleccionados):
            idx_abs = idx_inicio + idx_rel
            ax.plot(tiempo_ms[idx_abs], señal[idx_abs], 'ro' if tipo == 'pos' else 'bo')
            resultados_flicker.append({
                'Ojo': ojo_nombre,
                'Tipo': tipo,
                'Valor (µV)': f"{señal[idx_abs]:.2f}",
                'Tiempo (ms)': f"{tiempo_ms[idx_abs]:.1f}"
            })

        ax.plot(tiempo_ms, señal, label=nombre)
        ax.set_title(f"{ojo_nombre} - Flicker")
        ax.set_xlabel("Tiempo (ms)")
        ax.set_ylabel("µV")
        ax.legend()
        ax.grid()

    path_flicker_izq = "flicker_izq.png"
    path_flicker_der = "flicker_der.png"
    fig_flicker_izq.savefig(path_flicker_izq)
    fig_flicker_der.savefig(path_flicker_der)
    plt.close(fig_flicker_izq)
    plt.close(fig_flicker_der)

    row = doc.add_table(rows=1, cols=2).rows[0].cells
    row[0].paragraphs[0].add_run().add_picture(path_flicker_izq, width=Inches(3.0))
    row[1].paragraphs[0].add_run().add_picture(path_flicker_der, width=Inches(3.0))

    doc.add_heading("Tabla resumen de Flicker", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Ojo'
    hdr[1].text = 'Tipo'
    hdr[2].text = 'Valor (µV)'
    hdr[3].text = 'Tiempo (ms)'
    for res in resultados_flicker:
        row = table.add_row().cells
        row[0].text = res['Ojo']
        row[1].text = res['Tipo']
        row[2].text = res['Valor (µV)']
        row[3].text = res['Tiempo (ms)']


    # ====== Información prueba ======
    doc.add_heading("Información sobre la prueba", level=1)
    doc.add_paragraph(
        "Las intensidades de los estímulos usados durante las pruebas ERG se detallan a continuación, expresadas en cd·s·m⁻² y su correspondiente estímulo en mV:" 
    )

    tabla_info = doc.add_table(rows=1, cols=3)
    tabla_info.style = 'Table Grid'
    hdr = tabla_info.rows[0].cells
    hdr[0].text = 'Grupo'
    hdr[1].text = 'Intensidad (cd·s·m⁻²)'
    hdr[2].text = 'Estimulación (mV)'

    datos_intensidad = {
        1: ("0,0001", "2,45"),
        2: ("0,009", "2,55"),
        3: ("0,1013", "2,65"),
        4: ("0,9912", "3,25"),
        5: ("3,1724", "5"),
        6: ("3,1724", "5"),
        7: ("0,9912", "3,25"),
        8: ("3,1724", "5"),
        9: ("3,1724", "5")
    }

    for grupo, (intensidad, estimulo) in datos_intensidad.items():
        row = tabla_info.add_row().cells
        row[0].text = grupo_nombres[grupo]
        row[1].text = f"{intensidad} cd·s·m⁻²"
        row[2].text = f"{estimulo} mV"
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output




def mostrar_graficas_ojo_izquierdo(data, grupo_nombres, UMBRAL_TIEMPO=0.1):
    import matplotlib.pyplot as plt
    import numpy as np
    from scipy.signal import find_peaks
    import streamlit as st

    grupos_con_onda_a = [4, 5, 7, 8]
    grupos = sorted(data['Grupo'].unique())
    for grupo in grupos:
        nombre = grupo_nombres[grupo]
        subgrupo = data[data['Grupo'] == grupo]
        tiempo = subgrupo['Tiempo'].values
        señal = subgrupo['OjoIzquierdo'].values
        tiempo_ms = tiempo * 1000
        idx_inicio = np.argmax(tiempo >= UMBRAL_TIEMPO)

        fig, ax = plt.subplots(figsize=(18, 6))

        if grupo == 6:
            señal_post100 = señal[idx_inicio:]
            picos, _ = find_peaks(señal_post100)
            if len(picos) < 2:
                continue
            top2_indices = np.argsort(señal_post100[picos])[-2:]
            top2_sorted = top2_indices[np.argsort(señal_post100[picos][top2_indices])]
            segundo_pico_rel = picos[top2_sorted[0]]
            segundo_pico_idx = idx_inicio + segundo_pico_rel
            pre_pico = señal[idx_inicio:segundo_pico_idx]
            min_idx_rel = np.argmin(pre_pico)
            inicio_idx = idx_inicio + min_idx_rel
            ax.plot(tiempo_ms, señal, label='Señal ojo izquierdo')
            ax.plot(tiempo_ms[inicio_idx], señal[inicio_idx], 'mo', label='Inicio real')
            ax.plot(tiempo_ms[segundo_pico_idx], señal[segundo_pico_idx], 'ro', label='Segundo pico')

        elif grupo == 9:
            idx_inicio = np.argmax(tiempo >= 0.25)
            señal_post = señal[idx_inicio:]
            tiempo_post = tiempo_ms[idx_inicio:]
            picos_pos, _ = find_peaks(señal_post, prominence=20)
            picos_neg, _ = find_peaks(-señal_post, prominence=20)
            picos_total = sorted([(i, 'neg') for i in picos_neg] + [(i, 'pos') for i in picos_pos], key=lambda x: x[0])
            orden_deseado = ['neg', 'pos', 'neg', 'pos']
            seleccionados = []
            for idx, tipo in picos_total:
                if len(seleccionados) < len(orden_deseado) and tipo == orden_deseado[len(seleccionados)]:
                    seleccionados.append(idx_inicio + idx)
                if len(seleccionados) == 4:
                    break
            if len(seleccionados) == 4:
                min1, max1, min2, max2 = seleccionados
                ax.plot(tiempo_ms, señal, label='Señal ojo izquierdo')
                ax.axvline(250, linestyle='--', color='gray', label='Inicio 250ms')
                ax.plot(tiempo_ms[min1], señal[min1], 'bo', label='Pico negativo 1')
                ax.plot(tiempo_ms[max1], señal[max1], 'ro', label='Pico positivo 1')
                ax.plot(tiempo_ms[min2], señal[min2], 'bv', label='Pico negativo 2')
                ax.plot(tiempo_ms[max2], señal[max2], 'r^', label='Pico positivo 2')

        else:
            señal_post = señal[idx_inicio:]
            pico_b_rel = np.argmax(señal_post)
            pico_b_idx = idx_inicio + pico_b_rel
            pico_b_valor = señal[pico_b_idx]
            ax.plot(tiempo_ms, señal, label='Señal ojo izquierdo')
            ax.axvline(tiempo_ms[idx_inicio], linestyle='--', color='gray', label='Inicio 100ms')
            if grupo in grupos_con_onda_a:
                señal_a = señal[idx_inicio:pico_b_idx]
                pico_a_rel = np.argmin(señal_a)
                pico_a_idx = idx_inicio + pico_a_rel
                inicio_b_idx = pico_a_idx
                ax.plot(tiempo_ms[idx_inicio], señal[idx_inicio], 'ko', label='Inicio onda a')
                ax.plot(tiempo_ms[pico_a_idx], señal[pico_a_idx], 'go', label='Pico onda a')
            else:
                señal_inicial = señal[idx_inicio:pico_b_idx]
                min_rel = np.argmin(señal_inicial)
                inicio_b_idx = idx_inicio + min_rel
            ax.plot(tiempo_ms[inicio_b_idx], señal[inicio_b_idx], 'mo', label='Inicio onda b')
            ax.plot(tiempo_ms[pico_b_idx], señal[pico_b_idx], 'ro', label='Pico onda b')

        ax.set_title(nombre)
        ax.set_xlabel("Tiempo (ms)")
        ax.set_ylabel("µV")
        ax.legend()
        ax.grid()
        st.pyplot(fig)
